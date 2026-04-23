"""
Modulo Word - Agente Contabil
Geracao e preenchimento automatico de documentos .docx
com dados contabeis para o Departamento de Processos.
"""

import os
import logging
from typing import Dict, Optional
from datetime import datetime

logger = logging.getLogger(__name__)


class WordModule:
    """
    Modulo para criacao e manipulacao de documentos Word (.docx).
    Gera contratos, relatorios, declaracoes e cartas automaticamente.
    """

    def __init__(self):
        self.logger = logging.getLogger(self.__class__.__name__)
        self.templates_dir = 'dados/templates'
        os.makedirs(self.templates_dir, exist_ok=True)

    def preencher_template(self, template_nome: str, dados: Dict) -> Optional[str]:
        """
        Preenche um template Word com dados fornecidos.
        Substitui marcadores {{campo}} pelos valores.

        Args:
            template_nome: Nome do arquivo template
            dados: Dicionario com os dados para substituicao

        Returns:
            Caminho do arquivo gerado ou None em caso de erro
        """
        from docx import Document
        from docx.shared import Pt

        template_path = os.path.join(self.templates_dir, template_nome)
        if not os.path.exists(template_path):
            self.logger.warning(f"Template nao encontrado: {template_path}. Criando documento em branco.")
            return self.criar_documento_basico(dados)

        try:
            doc = Document(template_path)

            # Substitui nos paragrafos
            for para in doc.paragraphs:
                for chave, valor in dados.items():
                    marcador = f"{{{{{chave}}}}}"
                    if marcador in para.text:
                        for run in para.runs:
                            if marcador in run.text:
                                run.text = run.text.replace(marcador, str(valor))

            # Substitui nas tabelas
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for para in celula.paragraphs:
                            for chave, valor in dados.items():
                                marcador = f"{{{{{chave}}}}}"
                                if marcador in para.text:
                                    for run in para.runs:
                                        if marcador in run.text:
                                            run.text = run.text.replace(marcador, str(valor))

            # Salva arquivo gerado
            nome_saida = f"{template_nome.replace('.docx', '')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            saida_path = os.path.join('dados/relatorios', nome_saida)
            os.makedirs('dados/relatorios', exist_ok=True)
            doc.save(saida_path)
            self.logger.info(f"Documento gerado: {saida_path}")
            return saida_path

        except Exception as e:
            self.logger.error(f"Erro ao preencher template {template_nome}: {e}")
            return None

    def criar_documento_basico(self, dados: Dict) -> str:
        """
        Cria documento Word basico com os dados fornecidos.
        Usado quando nao ha template disponivel.
        """
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Cabecalho
        titulo = doc.add_heading('DOCUMENTO CONTABIL', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Data de emissao: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph('')

        # Adiciona os dados como tabela
        if dados:
            tabela = doc.add_table(rows=1, cols=2)
            tabela.style = 'Table Grid'

            # Cabecalho da tabela
            hdr = tabela.rows[0].cells
            hdr[0].text = 'Campo'
            hdr[1].text = 'Valor'

            for chave, valor in dados.items():
                row = tabela.add_row().cells
                row[0].text = str(chave)
                row[1].text = str(valor)

        doc.add_paragraph('')
        doc.add_paragraph('Agente Contabil - Departamento de Processos')

        saida_path = f"dados/relatorios/documento_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        os.makedirs('dados/relatorios', exist_ok=True)
        doc.save(saida_path)
        self.logger.info(f"Documento basico criado: {saida_path}")
        return saida_path

    def gerar_relatorio(self, dados_relatorio: Dict, tipo: str, caminho_saida: str) -> str:
        """
        Gera relatorio contabil formatado em Word.

        Args:
            dados_relatorio: Dicionario com dados do relatorio
            tipo: Tipo (DRE, balancete, bp, fluxo_caixa)
            caminho_saida: Caminho de destino do arquivo

        Returns:
            Caminho do arquivo gerado
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        titulos = {
            'DRE': 'DEMONSTRACAO DO RESULTADO DO EXERCICIO',
            'balancete': 'BALANCETE DE VERIFICACAO',
            'bp': 'BALANCO PATRIMONIAL',
            'fluxo_caixa': 'DEMONSTRACAO DO FLUXO DE CAIXA'
        }

        titulo_texto = titulos.get(tipo, tipo.upper())
        titulo = doc.add_heading(titulo_texto, 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        empresa = dados_relatorio.get('empresa', 'Empresa')
        periodo = dados_relatorio.get('periodo', datetime.now().strftime('%B/%Y'))

        sub = doc.add_paragraph(f"{empresa} | Periodo: {periodo}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')

        # Adiciona secoes do relatorio
        secoes = dados_relatorio.get('secoes', [])
        for secao in secoes:
            doc.add_heading(secao.get('titulo', ''), level=2)
            itens = secao.get('itens', [])
            if itens:
                tabela = doc.add_table(rows=1, cols=2)
                tabela.style = 'Table Grid'
                tabela.rows[0].cells[0].text = 'Descricao'
                tabela.rows[0].cells[1].text = 'Valor (R$)'
                for item in itens:
                    row = tabela.add_row().cells
                    row[0].text = item.get('descricao', '')
                    row[1].text = f"{item.get('valor', 0):,.2f}"
            doc.add_paragraph('')

        doc.add_paragraph(f"Documento gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph('Agente Contabil - Departamento de Processos')

        os.makedirs(os.path.dirname(caminho_saida) if os.path.dirname(caminho_saida) else '.', exist_ok=True)
        doc.save(caminho_saida)
        self.logger.info(f"Relatorio Word gerado: {caminho_saida}")
        return caminho_saida

    def gerar_declaracao(self, dados: Dict, tipo_declaracao: str) -> str:
        """
        Gera declaracoes contabeis padronizadas.
        Tipos: regularidade_fiscal, nada_consta, capacidade_financeira,
               enquadramento_me_epp, isencao_ir
        """
        templates_declaracoes = {
            'regularidade_fiscal': self._template_regularidade_fiscal,
            'nada_consta': self._template_nada_consta,
            'capacidade_financeira': self._template_capacidade_financeira,
        }

        gerador = templates_declaracoes.get(tipo_declaracao)
        if gerador:
            return gerador(dados)
        else:
            return self.criar_documento_basico(dados)

    def _template_regularidade_fiscal(self, dados: Dict) -> str:
        """Gera declaracao de regularidade fiscal"""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        titulo = doc.add_heading('DECLARACAO DE REGULARIDADE FISCAL', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('')
        texto = (
            f"Declaramos, para os devidos fins, que a empresa "
            f"{dados.get('empresa', '___________')}, "
            f"CNPJ n. {dados.get('cnpj', '________________')}, "
            f"com sede na {dados.get('endereco', '___________')}, "
            f"encontra-se em situacao de regularidade perante os orgaos fiscais federais, "
            f"estaduais e municipais na data de emissao deste documento."
        )
        doc.add_paragraph(texto)
        doc.add_paragraph('')
        doc.add_paragraph(
            f"{dados.get('cidade', 'Cidade')}, "
            f"{datetime.now().strftime('%d de %B de %Y')}."
        )
        doc.add_paragraph('')
        doc.add_paragraph('_' * 40)
        doc.add_paragraph(f"{dados.get('responsavel', 'Responsavel Tecnico')}")
        doc.add_paragraph(f"CRC: {dados.get('crc', '___________')}")

        saida = f"dados/relatorios/declaracao_regularidade_{datetime.now().strftime('%Y%m%d')}.docx"
        os.makedirs('dados/relatorios', exist_ok=True)
        doc.save(saida)
        return saida

    def _template_nada_consta(self, dados: Dict) -> str:
        """Gera certidao de nada consta"""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        titulo = doc.add_heading('CERTIDAO DE NADA CONSTA', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('')
        doc.add_paragraph(
            f"Certificamos que nada consta em nossos registros contra a empresa "
            f"{dados.get('empresa', '___________')}, "
            f"CNPJ {dados.get('cnpj', '________________')}, "
            f"ate a presente data."
        )
        doc.add_paragraph('')
        doc.add_paragraph(f"{dados.get('cidade', 'Cidade')}, {datetime.now().strftime('%d/%m/%Y')}.")

        saida = f"dados/relatorios/certidao_nada_consta_{datetime.now().strftime('%Y%m%d')}.docx"
        os.makedirs('dados/relatorios', exist_ok=True)
        doc.save(saida)
        return saida

    def _template_capacidade_financeira(self, dados: Dict) -> str:
        """Gera declaracao de capacidade financeira"""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        titulo = doc.add_heading('DECLARACAO DE CAPACIDADE FINANCEIRA', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('')
        doc.add_paragraph(
            f"Declaramos que a empresa {dados.get('empresa', '___________')}, "
            f"CNPJ {dados.get('cnpj', '________________')}, "
            f"possui capacidade financeira para honrar seus compromissos, "
            f"apresentando patrimonio liquido de R$ {dados.get('patrimonio', '0,00')} "
            f"conforme balanco patrimonial de {dados.get('data_balanco', '___/___/______')}."
        )

        saida = f"dados/relatorios/capacidade_financeira_{datetime.now().strftime('%Y%m%d')}.docx"
        os.makedirs('dados/relatorios', exist_ok=True)
        doc.save(saida)
        return saida
